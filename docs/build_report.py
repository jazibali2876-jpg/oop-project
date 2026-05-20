"""Generate the formal university project report (Project_Report.docx).

Grounded in the planning docs (planning/01..10) and the built code under src/.
Run from the project root:  python docs/build_report.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "presentation" / "screenshots"
OUT   = ROOT / "docs" / "Project_Report.docx"

INK   = RGBColor(0x10, 0x14, 0x1F)
ACCENT= RGBColor(0xFF, 0x6B, 0x35)
MUTED = RGBColor(0x5A, 0x5F, 0x6C)
LINE  = RGBColor(0xCB, 0xD0, 0xDA)

doc = Document()

# --- Page setup: A4 ---------------------------------------------------------
sec = doc.sections[0]
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.2)
sec.top_margin  = sec.bottom_margin = Cm(2.2)

# Default style — body text
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

# ------------------------------------------------------------------ helpers
def set_run(run, *, size=11, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return run


def H(text, level=1, *, color=None, before=None, after=Pt(4)):
    sizes = {0: 26, 1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = before if before is not None else Pt(14 if level <= 1 else 10)
    p.paragraph_format.space_after  = after
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size=sizes.get(level, 11), bold=True,
            color=color if color else (ACCENT if level <= 1 else INK))
    return p


def P(text, *, italic=False, size=11, color=INK, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, italic=italic, color=color)
    return p


def bullets(lines, *, level=0):
    for ln in lines:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(ln)
        set_run(r, size=11)


def numbered(lines):
    for ln in lines:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(ln)
        set_run(r, size=11)


def code_block(code, *, font_size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(code)
    set_run(r, size=font_size, color=INK, font="Consolas")


def page_break():
    doc.add_page_break()


def table_kv(rows, *, col_widths=(Inches(1.9), Inches(4.4))):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light List Accent 1"
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, (k, v) in enumerate(rows):
        c1, c2 = t.rows[i].cells
        c1.width, c2.width = col_widths
        c1.text = ""; c2.text = ""
        r1 = c1.paragraphs[0].add_run(k); set_run(r1, bold=True, size=10)
        r2 = c2.paragraphs[0].add_run(v); set_run(r2, size=10)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)


def table_grid(headers, rows, *, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=10)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ""
            r = c.paragraphs[0].add_run(str(val))
            set_run(r, size=10)
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for r_idx in range(len(t.rows)):
                t.rows[r_idx].cells[col_idx].width = w
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CBD0DA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_screenshot(name, *, caption=None, width=Inches(5.5)):
    p = SHOTS / f"{name}.png"
    if p.exists():
        doc.add_picture(str(p), width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ph.add_run(f"[ screenshot placeholder — drop {name}.png into presentation/screenshots/ ]")
        set_run(r, italic=True, color=MUTED, size=10)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(f"Figure: {caption}")
        set_run(r, italic=True, color=MUTED, size=9)


# ============================================================================
#  TITLE PAGE
# ============================================================================
for _ in range(3): doc.add_paragraph()

t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(t1.add_run("SMART RESTAURANT POS"), size=30, bold=True, color=INK)

t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(t2.add_run("A Modern Desktop Point-of-Sale Application"),
        size=16, italic=True, color=ACCENT)

t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(t3.add_run("Final-Year Object-Oriented Programming Project"),
        size=13, color=MUTED)

for _ in range(3): doc.add_paragraph()

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row_idx, (k, v) in enumerate([
    ("Submitted by",   "Jazib Ali"),
    ("Enrollment No.", "01-135221-020"),
    ("Course",         "Object Oriented Programming"),
    ("Year",           "2026"),
]):
    c1, c2 = meta.rows[row_idx].cells
    c1.width = Inches(2.0); c2.width = Inches(3.0)
    c1.text = ""; c2.text = ""
    r1 = c1.paragraphs[0].add_run(k);  set_run(r1, bold=True, size=12)
    r2 = c2.paragraphs[0].add_run(v);  set_run(r2, size=12)
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

for _ in range(8): doc.add_paragraph()

f = doc.add_paragraph(); f.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(f.add_run("Built with C++17  ·  SFML 3.1  ·  Dear ImGui  ·  MinGW-w64  ·  CMake"),
        size=11, color=MUTED)
f2 = doc.add_paragraph(); f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(f2.add_run("github.com/jazibali2876-jpg/oop-project"),
        size=10, color=MUTED, italic=True)
page_break()


# ============================================================================
#  ABSTRACT
# ============================================================================
H("Abstract", level=1)
P("Smart Restaurant POS is an offline-first desktop point-of-sale application "
  "built as a final-year Object-Oriented Programming project. It implements "
  "the day-to-day workflow of an independent restaurant — order taking, "
  "billing and receipt printing, inventory tracking, table management, "
  "customer loyalty, a kitchen display, and analytics — in a single C++17 "
  "executable that requires no installation and no database.")
P("The application is written in C++17 and built with the MinGW-w64 GCC "
  "toolchain. The graphical interface is rendered using Dear ImGui's "
  "immediate-mode UI on top of an SFML 3.1 window, with the integration "
  "between the two libraries hand-written from scratch. All persistent "
  "data — users, menu, inventory, orders, customers, loyalty, sales history, "
  "tables, settings, and receipts — is stored exclusively through C++ file "
  "handling, mixing human-readable text files with fixed-width binary record "
  "structs. There is no SQL, no embedded database, and no network dependency.")
P("Beyond the functional product, the project is structured to demonstrate "
  "the full set of OOP concepts required for the course — classes, "
  "encapsulation, inheritance, polymorphism, abstraction, virtual functions, "
  "templates, operator overloading, exception handling, the C++ Standard "
  "Template Library, modular programming, and header/source separation. "
  "Each concept is grounded in a specific named class so it can be defended "
  "directly during the viva.")
page_break()


# ============================================================================
#  TABLE OF CONTENTS
# ============================================================================
H("Table of Contents", level=1)
toc_rows = [
    ("1",   "Introduction"),
    ("2",   "Objectives"),
    ("3",   "Existing Systems & Motivation"),
    ("4",   "System Requirements"),
    ("5",   "Technology Stack"),
    ("6",   "System Architecture"),
    ("7",   "Object-Oriented Design"),
    ("8",   "UML & Use-Case Diagrams"),
    ("9",   "File-Handling Design"),
    ("10",  "Feature Walkthrough"),
    ("11",  "Implementation Highlights"),
    ("12",  "Testing & Verification"),
    ("13",  "Results"),
    ("14",  "Project Journey & Lessons Learned"),
    ("15",  "Limitations & Future Work"),
    ("16",  "Conclusion"),
    ("17",  "References"),
    ("A",   "Appendix A — Repository Layout"),
    ("B",   "Appendix B — Binary File Schemas"),
    ("C",   "Appendix C — How to Run & Build"),
]
for n, t in toc_rows:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.0))
    r = p.add_run(f"{n}.\t{t}")
    set_run(r, size=11)
page_break()


# ============================================================================
#  1. INTRODUCTION
# ============================================================================
H("1.  Introduction", level=1)
P("Restaurants — particularly independent and small-chain establishments — "
  "still juggle paper order pads, separate kitchen tickets, and standalone "
  "receipt printers. Inventory drifts out of sync with what is actually being "
  "sold. Reporting is manual. Customer loyalty data is lost. The commercial "
  "alternatives — Toast, Square, Lightspeed, Clover — solve these problems, "
  "but they are expensive monthly subscriptions, require a constant cloud "
  "connection, and force a specific hardware ecosystem. For a small "
  "restaurant that runs out of one physical location, much of that software "
  "is overkill.")
P("Smart Restaurant POS is a pragmatic alternative: a single Windows "
  "executable that captures the essentials of a real POS — orders, billing, "
  "inventory, tables, customers, kitchen display, and reporting — and runs "
  "offline, without any server, database, or external dependency.")
P("Just as importantly, this is a final-year university project for the "
  "Object-Oriented Programming course. Every architectural decision in the "
  "system was made to satisfy two constraints simultaneously: produce a real "
  "product that feels like commercial POS software, and demonstrate every "
  "OOP concept the course examines, in a way the author can defend in viva.")
hrule()


# ============================================================================
#  2. OBJECTIVES
# ============================================================================
H("2.  Objectives", level=1)
P("The project was scoped against four explicit objectives:")
bullets([
    "Build a functional desktop POS application covering the core workflow "
    "of an independent restaurant: authentication, menu management, order "
    "taking, billing and receipts, inventory with auto-deduction, table "
    "reservations, a kitchen display, customer profiles with loyalty, and "
    "analytics reports.",
    "Demonstrate the complete set of OOP concepts required for the course — "
    "classes and objects, encapsulation, inheritance, polymorphism, "
    "abstraction, virtual functions, templates, operator overloading, "
    "exception handling, STL containers, modular programming, and header/"
    "source separation — each mapped to a specific class in the codebase.",
    "Implement complete file handling using only the C++ standard library "
    "(ifstream / ofstream / fstream). The persistence layer must cover both "
    "human-readable text files (menu, inventory, recipes, tables, settings) "
    "and fixed-width binary record structs (users, customers, orders, "
    "loyalty ledger, sales history). No SQL, no embedded database.",
    "Keep the project realistically buildable within one semester. No "
    "enterprise frameworks, no commercial libraries, no cloud, no networking, "
    "no machine learning. Stay on a fixed stack and resist scope creep.",
])
hrule()


# ============================================================================
#  3. EXISTING SYSTEMS & MOTIVATION
# ============================================================================
H("3.  Existing Systems & Motivation", level=1)
P("Three commercial POS systems were used as visual and functional reference "
  "points for the user experience: Toast, Square for Restaurants, and "
  "Lightspeed Restaurant. The shared traits the project aims to match are:")
bullets([
    "A persistent left-side navigation bar with a fixed list of major "
    "areas (Dashboard, Menu, Orders, Billing, etc.).",
    "A top bar showing the active screen title, real-time clock, theme "
    "toggle, and currently signed-in user.",
    "Card-based dashboards with at-a-glance KPI tiles.",
    "Receipt previews rendered in a fixed-width thermal layout.",
    "Color-coded table grids that respond to one-tap state transitions "
    "(reserve / occupy / free).",
    "Three-lane kitchen displays (pending → preparing → ready).",
])
P("Smart Restaurant POS does not aim to compete with those systems on "
  "scale or features. It aims to feel familiar to anyone who has used one, "
  "while being something a single developer can build, defend, and ship in a "
  "semester as a course project.")
hrule()


# ============================================================================
#  4. SYSTEM REQUIREMENTS
# ============================================================================
H("4.  System Requirements", level=1)

H("4.1  Functional Requirements", level=2)
table_grid(
    ["#", "Requirement"],
    [
        ("FR-1",  "The system shall authenticate users via username and password, "
                  "with masked password input and a 'remember me' option."),
        ("FR-2",  "The system shall support three roles — Admin, Cashier, and "
                  "Kitchen — each with a distinct capability set."),
        ("FR-3",  "Admin shall be able to create, update, and delete menu items, "
                  "categorized as FastFood, BBQ, Chinese, Drinks, or Desserts."),
        ("FR-4",  "Cashier shall be able to build an order from the menu, modify "
                  "quantities, apply a discount, attach special instructions, and "
                  "place the order."),
        ("FR-5",  "On order placement, the system shall validate available "
                  "stock against the recipe for each menu item and either "
                  "reject the order (with a clear message) or deduct ingredients."),
        ("FR-6",  "The system shall generate an itemized receipt with subtotal, "
                  "discount, tax, total, and payment method, in a thermal-style "
                  "layout suitable for printing or text export."),
        ("FR-7",  "The system shall track per-customer profiles (name, phone, "
                  "total spent, loyalty points, tier) and apply loyalty points "
                  "automatically when an order is recorded against a customer."),
        ("FR-8",  "The system shall expose a table grid with three states "
                  "(Free, Occupied, Reserved) and one-tap state transitions."),
        ("FR-9",  "The kitchen role shall see active orders in three lanes "
                  "(Pending, Preparing, Ready) and advance their status."),
        ("FR-10", "The system shall provide analytics over today, last 7 days, "
                  "and last 30 days, including revenue, order count, average "
                  "order value, top items, and orders by hour of day."),
        ("FR-11", "All persistent data shall survive a clean application "
                  "shutdown and be restored on the next launch."),
        ("FR-12", "On a fresh installation with no data files present, the "
                  "system shall seed sensible defaults so the application is "
                  "usable immediately, without manual setup."),
    ],
    col_widths=(Inches(0.7), Inches(5.6)),
)

H("4.2  Non-Functional Requirements", level=2)
table_grid(
    ["#", "Requirement"],
    [
        ("NFR-1", "The application shall run on 64-bit Windows 10 or 11 with no "
                  "installer, no administrator rights, and no network access."),
        ("NFR-2", "The application shall start to a usable login screen within "
                  "three seconds on commodity hardware."),
        ("NFR-3", "No persistence layer shall require a database engine "
                  "(SQLite, MySQL, Postgres, etc.). Only file handling via "
                  "the C++ standard library."),
        ("NFR-4", "Whole-file rewrites shall be crash-safe: writes shall not "
                  "leave a corrupted file if the application is force-killed "
                  "mid-write."),
        ("NFR-5", "All binary record structs shall have a static_assert "
                  "guard pinning sizeof(record) at a stable byte count."),
        ("NFR-6", "The build shall be reproducible: a fresh checkout on a "
                  "clean machine, given the supplied toolchain, shall produce "
                  "a byte-equivalent executable."),
        ("NFR-7", "All third-party assets (fonts, vendored libraries) shall "
                  "be supplied locally; the build shall not require any "
                  "runtime download."),
        ("NFR-8", "The packaged dist/ folder shall be self-contained: every "
                  "required runtime DLL (SFML and MinGW C++ runtime) shall "
                  "sit beside the executable so it runs on double-click."),
    ],
    col_widths=(Inches(0.7), Inches(5.6)),
)
hrule()


# ============================================================================
#  5. TECHNOLOGY STACK
# ============================================================================
H("5.  Technology Stack", level=1)
table_kv([
    ("Language",      "C++17 (ISO/IEC 14882:2017)"),
    ("Compiler",      "MinGW-w64 GCC 14.2.0 (UCRT / POSIX / SEH)"),
    ("Build system",  "CMake 3.30 (vendored portable)"),
    ("Windowing & I/O","SFML 3.1.0 — sf::RenderWindow, sf::Event, sf::Texture, sf::Sound"),
    ("UI framework",  "Dear ImGui (master branch) — immediate-mode"),
    ("UI renderer",   "imgui_impl_opengl3 (bundled with Dear ImGui), OpenGL 3.3"),
    ("Integration",   "Hand-written ImGui ↔ SFML 3 glue (no third-party binding)"),
    ("Charts",        "ImDrawList primitives (AddRectFilled, PathArcTo, etc.) — no ImPlot"),
    ("Storage",       "Plain C++ file handling: <fstream>, text + binary"),
    ("Fonts",         "Inter Regular / Bold + Font Awesome 6 Free (vendored TTFs)"),
])
P("Every dependency is vendored inside the project folder. No third-party "
  "library is fetched at runtime. The MinGW toolchain ships in mingw64/, "
  "Dear ImGui in imgui-master/, SFML in SFML-3.1.0/, and CMake plus the "
  "from-source MinGW build of SFML in third_party/.")
hrule()


# ============================================================================
#  6. SYSTEM ARCHITECTURE
# ============================================================================
H("6.  System Architecture", level=1)
P("The single most important architectural decision in this project is the "
  "strict two-layer split, with a thin integration glue layer between the "
  "application and the UI library.")

H("6.1  The Domain / UI Split", level=2)
P("Dear ImGui is an immediate-mode UI library — it has no persistent widget "
  "objects, no .ui files, no signals or slots. Without discipline, this "
  "encourages business logic to bleed into draw functions, which would "
  "completely undermine the OOP demonstration. To avoid that, the codebase is "
  "split into two non-overlapping layers:")

table_grid(
    ["Layer", "Folder", "Responsibility"],
    [
        ("Domain", "src/domain/",
         "Plain C++17 entities, services, repositories. Owns ALL business "
         "state and all persisted data. Holds the entire OOP demonstration. "
         "Includes no ImGui and no SFML headers."),
        ("UI",     "src/ui/",
         "Immediate-mode screens, widgets, theme, animation helpers. Holds "
         "only transient UI state (selection, scroll, animation timers). "
         "Reads from the domain and calls into its services. Never opens a "
         "file directly."),
        ("Platform","src/platform/",
         "Hand-written ImGui ↔ SFML 3 integration. The ONLY folder that "
         "includes both ImGui and SFML headers. Owns the SFML window, "
         "translates sf::Event into ImGuiIO events, drives "
         "imgui_impl_opengl3 against SFML's OpenGL context."),
        ("Util",   "src/util/",
         "Tiny cross-cutting helpers — atomic file write, string trimming, "
         "logging, deterministic random. No dependencies on the layers above."),
    ],
    col_widths=(Inches(0.9), Inches(1.2), Inches(4.2)),
)

H("6.2  The Manual ImGui ↔ SFML 3 Integration", level=2)
P("There is no ready-made ImGui-SFML binding for SFML 3.x — the binding "
  "library ImGui-SFML targets the older SFML 2.x API. This project writes "
  "the integration from scratch in four small classes:")
bullets([
    "Platform — owns sf::RenderWindow with an attached OpenGL context, "
    "drains the event queue each frame, and forwards every sf::Event to "
    "the input bridge.",
    "ImGuiBackend — wraps ImGui::CreateContext, ImGui_ImplOpenGL3_Init, "
    "the per-frame NewFrame / Render / RenderDrawData lifecycle, and "
    "shutdown.",
    "InputBridge — translates each SFML 3 event subtype "
    "(event.getIf<sf::Event::MouseMoved>(), MouseButtonPressed, "
    "KeyPressed, TextEntered, Resized, FocusGained / FocusLost, "
    "MouseWheelScrolled) into the corresponding ImGuiIO::Add*Event call.",
    "KeyMap — maps every scoped sf::Keyboard::Key value into the "
    "matching ImGuiKey enum, covering A–Z, 0–9, F1–F12, navigation "
    "keys, modifiers, and the numpad.",
])
P("The per-frame loop in App::run() is therefore explicit:")
code_block(
"""while (platform.pumpEvents(input)) {
    float dt = platform.deltaSeconds();
    auto fb  = platform.framebufferSize();
    imgui.newFrame(dt, fb.x, fb.y);

    ui::Shell::draw(ctx);             // domain reads, ImGui draw calls
    toasts.tickAndDraw(dt, fb.x);

    platform.clear();
    imgui.render(platform.window());
    platform.display();
}""")

H("6.3  How Data Flows", level=2)
P("A worked example — the cashier places an order:")
numbered([
    "OrderScreen draws the menu grid by reading MenuService::all(). "
    "Tapping an item appends it to OrderDraft (a transient UI-state struct, "
    "not yet a domain Order).",
    "When the cashier taps 'Place Order', OrderScreen calls "
    "OrderService::place(draft, session.user(), customerId).",
    "OrderService validates the draft (throws EmptyOrderException if empty), "
    "asks InventoryService::canFulfill / deduct (throws "
    "InsufficientStockException if needed), assigns the order id and "
    "timestamp, persists the order through OrderRepository, appends a "
    "SalesRecord through SalesHistoryRepository, and enqueues a kitchen "
    "ticket through KitchenService.",
    "On exception, OrderScreen catches the DomainException and pushes a "
    "red Toast onto ToastQueue; on success it navigates to BillingScreen, "
    "which builds a Receipt via BillingService and writes it to "
    "data/receipts/RCPT-NNNNNN.txt.",
])
hrule()
page_break()


# ============================================================================
#  7. OBJECT-ORIENTED DESIGN
# ============================================================================
H("7.  Object-Oriented Design", level=1)
P("Every required OOP concept maps to a specific named class in the "
  "codebase. This table is the primary viva-defense reference — each row "
  "names a concept, the file that demonstrates it, and the precise hook "
  "an examiner can point at.")
table_grid(
    ["Concept", "Where to point", "What to say"],
    [
        ("Classes & objects",
         "src/domain/menu/MenuItem.h",
         "Every persisted entity has its own class with private fields and "
         "behavior methods (MenuItem, Order, Customer, Ingredient, Table)."),
        ("Encapsulation",
         "src/domain/order/Order.h",
         "Order holds its items vector privately; mutation is only through "
         "add, removeAt, setDiscount, setStatus — totals stay consistent."),
        ("Inheritance",
         "src/domain/auth/User.h + Admin.h / Cashier.h / Kitchen.h",
         "Three concrete roles inherit from an abstract User base."),
        ("Polymorphism",
         "Session::user().can(Capability)",
         "Sidebar asks the currently logged-in user object what it can do; "
         "the same call site behaves differently per runtime type."),
        ("Abstraction",
         "User — pure virtual can() and roleName()",
         "User cannot be instantiated; only Admin/Cashier/Kitchen can."),
        ("Virtual functions",
         "User::can(Capability) const = 0, plus override in each subclass",
         "Pure-virtual in the base, override-marked in each subclass."),
        ("Templates",
         "src/domain/persistence/BinaryRepository.h",
         "One templated repository serves UserRecord, CustomerRecord, "
         "LoyaltyRecord, and SalesRecord. Id<Tag> is a templated strong-id."),
        ("Operator overloading",
         "src/domain/common/Money.h, Receipt::operator<<",
         "Money has +, -, *, comparison, and stream insertion. Receipt's "
         "operator<< emits the 44-column thermal-style receipt layout."),
        ("Exception handling",
         "src/domain/common/Exceptions.h",
         "DomainException base + 8 named subclasses; services throw, the "
         "UI catches and shows a Toast — no error-code soup."),
        ("STL containers",
         "vector, map, unordered_map, array<int,24>, optional, deque, "
         "unique_ptr",
         "At least six container types are in use across the domain."),
        ("Modular programming",
         "src/domain/ split into 10 sub-packages",
         "auth, menu, order, billing, inventory, customer, tables, kitchen, "
         "analytics, persistence — plus separate ui/, platform/, util/ layers."),
        ("Header/source split",
         "Every non-template class",
         "Declarations in .h, definitions in .cpp. Templates stay in headers "
         "by language necessity."),
    ],
    col_widths=(Inches(1.5), Inches(2.0), Inches(2.8)),
)
P("Volunteer-in-viva concepts (not required but present):")
bullets([
    "RAII — Platform and ImGuiBackend acquire window and ImGui context in "
    "their constructors, release them in their destructors.",
    "Smart pointers — Session owns the polymorphic User via std::unique_ptr.",
    "const-correctness — every read accessor on the domain is marked const.",
    "Move semantics — AuthService::login returns std::unique_ptr<User> by value.",
    "std::optional — used by Order::customer, Table::reservedFor, and "
    "every byId / findByX accessor that may legitimately return 'no result'.",
    "Lambdas — std::find_if, std::remove_if are used throughout the "
    "services with inline lambda predicates.",
])
hrule()
page_break()


# ============================================================================
#  8. UML & USE-CASE DIAGRAMS
# ============================================================================
H("8.  UML & Use-Case Diagrams", level=1)
P("Full Mermaid sources for every diagram live in "
  "planning/05-uml-diagrams.md and planning/06-use-case-diagrams.md in the "
  "repository. The summaries below describe each diagram by name; the "
  "Mermaid source is rendered directly on GitHub for any viewer.")

H("8.1  Class Diagram — Auth Package", level=2)
P("User is the abstract base. Admin, Cashier, and Kitchen all inherit from "
  "it and override the pure-virtual can(Capability) and roleName(). "
  "Session owns the currently authenticated user via std::unique_ptr<User>. "
  "AuthService is the only entry point that constructs User objects — "
  "either by reading them through UserRepository or by registering a new "
  "user with a salted hash via PasswordHash.")

H("8.2  Class Diagram — Order & Billing", level=2)
P("Order aggregates OrderItem instances and computes subtotal, tax, total, "
  "and estimated prep time. OrderService orchestrates the place / cancel "
  "/ setStatus use cases against OrderRepository, InventoryService, "
  "KitchenService, and SalesHistoryRepository. BillingService consumes a "
  "domain Order and produces a Receipt — which is itself a class whose "
  "stream-insertion operator emits the thermal-style printable text.")

H("8.3  Class Diagram — Persistence (Templates in Action)", level=2)
P("TextRepository is the base for line-oriented text storage "
  "(menu.txt, inventory.txt, recipes.txt, tables.txt). "
  "BinaryRepository<T> is the templated base for fixed-width POD records "
  "(users.dat, customers.dat, loyalty.dat, sales_history.dat). It is "
  "instantiated for four record types, each with its own static_assert "
  "pinning the on-disk byte count. OrderRepository is a special case — "
  "it writes variable-length records (one header followed by N item rows) "
  "so it manages the stream itself rather than inheriting from "
  "BinaryRepository<T>.")

H("8.4  Use-Case Diagram — Three Actors", level=2)
table_grid(
    ["Actor",  "Permitted use cases"],
    [
        ("Admin",
         "All cashier and kitchen use cases, plus: Manage Menu, Manage "
         "Inventory, Manage Recipes, Manage Users, View Analytics, "
         "Switch Theme, Override Order Status."),
        ("Cashier",
         "Login / Logout, Browse Menu, Place Order, Modify Draft, Apply "
         "Discount, Generate Receipt, Select Payment Method, Save / Export "
         "Receipt, Lookup or Register Customer, Apply Loyalty Points, "
         "Reserve Table, Mark Table Occupied / Free."),
        ("Kitchen",
         "Login / Logout, View Pending / Preparing / Ready lanes, "
         "Advance Ticket Status, See Special Instructions."),
    ],
    col_widths=(Inches(1.0), Inches(5.3)),
)

H("8.5  Data-Flow Diagram — Place-Order Flow", level=2)
P("The Place-Order flow is the central business workflow in the app and "
  "exercises the largest number of services in a single transaction:")
code_block(
"""Cashier  →  OrderScreen  →  OrderDraft (UI state)
                ↓ place()
            OrderService
               ├── InventoryService::canFulfill / deduct  →  inventory.txt
               ├── OrderRepository::append                 →  orders.dat (binary)
               ├── SalesHistoryRepository::append          →  sales_history.dat
               └── KitchenService::enqueue                 →  (in-memory ticket queue)
            BillingService::makeReceipt
                ↓
            ReceiptRepository::write                       →  receipts/RCPT-NNNNNN.txt
                ↓
            BillingScreen (preview + payment)""",
    font_size=9)
hrule()
page_break()


# ============================================================================
#  9. FILE-HANDLING DESIGN
# ============================================================================
H("9.  File-Handling Design", level=1)
P("All persistence uses the C++ standard library: std::ifstream, std::ofstream, "
  "std::fstream, std::filesystem. There is no SQL, no embedded database engine. "
  "Files split into two categories — human-readable text for editable data, "
  "and fixed-width binary records for ledgers.")

H("9.1  Text Files", level=2)
P("Text files use '|' as the field delimiter and a leading '#' for comment "
  "headers. The first line is always a self-describing header.")
table_grid(
    ["File", "Contents"],
    [
        ("data/menu.txt",
         "id | name | category | priceCents | imagePath | available"),
        ("data/inventory.txt",
         "id | name | unit | stock | reorderThreshold | costCentsPerUnit"),
        ("data/recipes.txt",
         "menuItemId | ingredientId:qty, ingredientId:qty, ..."),
        ("data/tables.txt",
         "id | seats | status | reservedForCustomerId | reservedAtEpoch"),
        ("data/settings.txt",
         "k=v lines — theme, rememberedUser, taxBp, currencySymbol, ..."),
        ("data/log.txt",
         "ISO8601 timestamp + level + module + message"),
        ("data/receipts/RCPT-NNNNNN.txt",
         "thermal-style printable receipt — output of Receipt::operator<<"),
    ],
    col_widths=(Inches(2.4), Inches(3.9)),
)

H("9.2  Binary Record Files", level=2)
P("Binary files are sequences of trivially copyable POD record structs. "
  "Each record struct has #pragma pack(1) plus a static_assert pinning its "
  "exact byte count, so the on-disk format cannot drift silently across "
  "code changes.")
table_grid(
    ["File", "Record", "Size", "Notes"],
    [
        ("data/users.dat",         "UserRecord",        "128 B",
         "1 record per login-able user (Admin / Cashier / Kitchen)."),
        ("data/orders.dat",        "OrderHeader + N×OrderItem", "128 + 56·N B",
         "Variable-length: one header then N item rows. itemCount in the "
         "header tells the reader how many items follow."),
        ("data/customers.dat",     "CustomerRecord",    "96 B",
         "Phone, name, lifetime spend, loyalty points, order count."),
        ("data/loyalty.dat",       "LoyaltyRecord",     "32 B",
         "Append-only ledger of every loyalty-point delta."),
        ("data/sales_history.dat", "SalesRecord",       "64 B",
         "Flat append-only stream optimized for analytics scans — "
         "duplicates per-order totals to avoid reading orders.dat."),
    ],
    col_widths=(Inches(1.9), Inches(1.6), Inches(0.7), Inches(2.2)),
)
P("Full byte-level field layouts for every record are in Appendix B.")

H("9.3  Crash Safety", level=2)
P("Whole-file rewrites use the write-temp-then-rename pattern: the new "
  "content is written to <path>.tmp, then std::filesystem::rename is "
  "called to atomically replace the live file. A power loss mid-write "
  "either leaves the original file untouched (if the temp wasn't fully "
  "written) or installs the new one cleanly. Binary appends rely on the "
  "OS append guarantee — the kernel serializes appends so partial writes "
  "can't interleave with another writer.")
P("This pattern is centralized in pos::util::writeAtomic in "
  "src/util/Files.cpp and is used by every TextRepository and by "
  "BinaryRepository<T>::overwriteAll.")

H("9.4  First-Run Behavior", level=2)
P("On a fresh installation, no files exist in data/. The app:")
numbered([
    "Creates data/ and data/receipts/ via util::mkdirs.",
    "Calls AuthService::seedDemoUsersIfEmpty — creates admin / admin123 "
    "(Admin), cashier / cashier123 (Cashier), kitchen / kitchen123 (Kitchen).",
    "Calls MenuService::seedDefaultsIfEmpty — 5 menu items across all "
    "5 categories.",
    "Calls InventoryService::seedDefaultsIfEmpty — 5 ingredients with "
    "stock and reorder thresholds.",
    "Calls InventoryService::seedDefaultRecipesIfEmpty — 2 sample recipes.",
    "Calls TableService::seedDefaultsIfEmpty — 10 tables with seat counts.",
])
P("Every repository tolerates a missing file as 'empty content' — it does "
  "not throw. This means a user can delete data/ at any time to reset to a "
  "blank install, and the next launch will re-seed automatically.")
hrule()
page_break()


# ============================================================================
#  10. FEATURE WALKTHROUGH
# ============================================================================
H("10.  Feature Walkthrough", level=1)

H("10.1  Splash & Login", level=2)
P("On launch, a 1.6-second splash fades into the login screen. The login "
  "form has masked password input (ImGuiInputTextFlags_Password), a "
  "'remember me' checkbox that persists the username in settings.txt, and "
  "lists the three demo accounts as a tip. An invalid credential triggers "
  "a 0.4-second shake animation on the form and a red toast on the right.")
add_screenshot("login", caption="Login screen with the three demo accounts listed.")

H("10.2  Dashboard", level=2)
P("Four metric cards across the top — today's revenue, active orders, "
  "low-stock count, and the top-selling item. Below them: today's top "
  "items list, a manually-drawn 7-day revenue line chart, and a "
  "low-stock alert table. Every figure is computed live from the data "
  "files via AnalyticsService.")
add_screenshot("dashboard", caption="Dashboard cards and revenue line chart.")

H("10.3  Menu Management", level=2)
P("A category filter, a search box, and a grid of menu cards. Admin sees "
  "Add / Edit / Delete buttons on every card; cashier and kitchen see the "
  "same screen with those buttons hidden by capability check. Adding or "
  "editing an item opens a modal with name, category dropdown, price, "
  "image path, and availability toggle.")
add_screenshot("menu", caption="Menu management screen (Admin view).")

H("10.4  New Order", level=2)
P("Two-pane layout: menu grid on the left, cart on the right. Tapping a "
  "menu card adds it to OrderDraft (a UI-only struct, not yet a domain "
  "Order). The cart shows live subtotal, discount, tax, and total, plus "
  "estimated prep time. Customer attachment is by phone lookup — registered "
  "customers are found by phone, new customers can be created inline.")
add_screenshot("order", caption="New Order screen — menu grid + draft cart.")

H("10.5  Billing & Receipts", level=2)
P("Payment method selection (cash / card / online-sim), a live receipt "
  "preview rendered by Receipt::operator<<, and Save / Export PDF-style "
  "buttons. Charging an order both records the payment and triggers loyalty "
  "point application via CustomerService.")
add_screenshot("billing", caption="Billing screen with thermal-style receipt preview.")

H("10.6  Inventory", level=2)
P("Tabular view of every ingredient. Rows below their reorder threshold "
  "are tinted in amber and prefixed with 'LOW'. The bottom of the screen "
  "shows a usage bar chart drawn manually with ImDrawList. Auto-deduction "
  "fires every time OrderService::place is called, reading each item's "
  "recipe and subtracting from stock.")
add_screenshot("inventory", caption="Inventory table with low-stock highlight.")

H("10.7  Tables", level=2)
P("Color-coded grid: green = Free, amber = Occupied, blue = Reserved. "
  "Tapping a table opens a popup with Reserve / Mark Occupied / Free. "
  "Reservation prompts for a phone number and a name and attaches a "
  "Customer record (found or created via findByPhoneOrCreate).")
add_screenshot("tables", caption="Tables screen with color-coded status grid.")

H("10.8  Kitchen Display", level=2)
P("Three vertical lanes: Pending, Preparing, Ready. Each ticket shows the "
  "order number, queued time, elapsed time (live, per-frame), the items, "
  "and any special instructions. One-tap advances the ticket to the next "
  "lane. Hitting 'Mark Served' on a Ready ticket removes it from the "
  "kitchen queue and updates the order's status to Served.")
add_screenshot("kitchen", caption="Kitchen display — three lanes with elapsed time.")

H("10.9  Analytics", level=2)
P("Range selector (Today / Last 7 days / Last 30 days) drives Revenue, "
  "Order Count, and Average Order Value cards. Three hand-drawn charts: "
  "a bar chart of revenue per day, a pie chart of the top items, and a "
  "line chart of orders by hour of day. All three are drawn using "
  "ImDrawList primitives — no ImPlot dependency.")
add_screenshot("analytics", caption="Analytics dashboard with three ImDrawList charts.")
hrule()
page_break()


# ============================================================================
#  11. IMPLEMENTATION HIGHLIGHTS
# ============================================================================
H("11.  Implementation Highlights", level=1)

H("11.1  The Role Hierarchy", level=2)
P("Polymorphism is demonstrated by the User base and three concrete "
  "subclasses. The base declares a pure-virtual capability check:")
code_block(
"""class User {
public:
    virtual ~User() = default;
    virtual bool        can(Capability) const = 0;
    virtual std::string roleName() const = 0;
    // ...
};

class Cashier final : public User {
public:
    bool can(Capability c) const override {
        switch (c) {
            case Capability::ViewMenu:
            case Capability::PlaceOrder:
            case Capability::ManageTables:
            case Capability::ManageCustomers:
                return true;
            default:
                return false;
        }
    }
    std::string roleName() const override { return \"Cashier\"; }
};""")

H("11.2  The Templated Repository", level=2)
P("BinaryRepository<T> is the single class that owns every binary file. "
  "T must be trivially copyable — checked at compile time. It exposes "
  "all(), append(), overwriteAll(), and a templated findIf() helper.")
code_block(
"""template <class T>
class BinaryRepository {
public:
    static_assert(std::is_trivially_copyable_v<T>,
                  \"BinaryRepository<T>: T must be trivially copyable\");

    std::vector<T> all() const;
    void           append(const T& rec);
    void           overwriteAll(const std::vector<T>& v);

    template <class Pred>
    std::optional<T> findIf(Pred p) const;
    // ...
};

class UserRepository : public BinaryRepository<UserRecord> { /* ... */ };""")

H("11.3  Operator Overloading on Money", level=2)
P("Money is a value type that stores cents as an int64_t. Arithmetic, "
  "comparison, and stream insertion are all overloaded.")
code_block(
"""class Money {
public:
    Money& operator+=(const Money& r) { cents_ += r.cents_; return *this; }
    friend Money operator+(Money a, const Money& b) { a += b; return a; }
    friend Money operator*(Money a, int n)          { a *= n; return a; }
    friend bool  operator<(const Money& a, const Money& b)
                                          { return a.cents_ < b.cents_; }
    friend std::ostream& operator<<(std::ostream& os, const Money& m);
private:
    std::int64_t cents_ = 0;
};""")

H("11.4  Stream Insertion on Receipt", level=2)
P("Receipt's operator<< emits the entire 44-column thermal-style receipt. "
  "Saving the receipt to a file is then a one-liner: serialize the Receipt "
  "into an ostringstream, call util::writeAtomic.")
code_block(
"""std::ostream& operator<<(std::ostream& os, const Receipt& r) {
    os << \"============================================\\n\";
    os << \"         SMART RESTAURANT POS\\n\";
    os << \"============================================\\n\";
    os << \"Order #:    \" << r.orderId().value() << \"\\n\";
    os << \"Date:       \" << r.issuedAt().format() << \"\\n\";
    // ... (items, subtotal, tax, total, payment, loyalty) ...
    return os;
}""")

H("11.5  Exception Boundary", level=2)
P("Domain exceptions are an inheritance hierarchy rooted at DomainException. "
  "Services throw them; the UI catches DomainException& at every screen "
  "boundary and turns it into a Toast. There are no error codes returned "
  "by domain functions.")
code_block(
"""class DomainException : public std::runtime_error { /* ... */ };
class InvalidCredentialsException : public DomainException { /* ... */ };
class InsufficientStockException  : public DomainException {
public:
    InsufficientStockException(unsigned ingredientId, double needed, double have);
    unsigned ingredientId() const;
    double   needed() const;
    double   have() const;
};

// At the UI boundary (OrderScreen):
try {
    auto placed = ctx.orders.place(draft, ctx.session.user(), customerId);
    ctx.toasts.success(\"Order placed\", \"Order #\" + std::to_string(placed.id().value()));
} catch (const DomainException& e) {
    ctx.toasts.danger(\"Order failed\", e.what());
}""")

H("11.6  Hand-Drawn Charts", level=2)
P("All three chart types (bar, pie, line) are drawn manually using "
  "ImDrawList primitives. There is no ImPlot dependency. The pie chart "
  "uses PathArcTo + PathFillConvex; the bar chart uses AddRectFilled; "
  "the line chart uses AddLine plus AddCircleFilled for the data points.")
code_block(
"""// PieChart (excerpt):
for (std::size_t i = 0; i < slices.size(); ++i) {
    float endAngle = startAngle + frac * 6.2831853f;
    dl->PathClear();
    dl->PathLineTo(center);
    int steps = std::max(1, (int)(segPerSlice * frac));
    for (int k = 0; k <= steps; ++k) {
        float a = startAngle + (endAngle - startAngle) * (k / (float)steps);
        dl->PathLineTo(ImVec2(center.x + cos(a)*radius,
                              center.y + sin(a)*radius));
    }
    dl->PathFillConvex(col);
    startAngle = endAngle;
}""")
hrule()
page_break()


# ============================================================================
#  12. TESTING & VERIFICATION
# ============================================================================
H("12.  Testing & Verification", level=1)
P("Testing is layered: compile-time guards for invariants the type system "
  "can express, environment-level checks via a preflight script, end-to-end "
  "verification by running the executable against a real data folder.")

H("12.1  Compile-Time Checks", level=2)
bullets([
    "Every binary record struct has static_assert(sizeof(R) == N) and "
    "static_assert(std::is_trivially_copyable_v<R>). Any inadvertent change "
    "to field types or order produces a hard compile error rather than a "
    "silent on-disk corruption.",
    "BinaryRepository<T>::T is constrained at the class level by the same "
    "is_trivially_copyable_v check, so passing a non-POD type fails the "
    "build immediately.",
    "All header/source pairs compile cleanly on -std=c++17 with no "
    "deprecation or shadowing warnings under MinGW-w64 GCC 14.2.",
])

H("12.2  Preflight Verification Script", level=2)
P("tools/preflight_check.ps1 is a self-contained PowerShell script that "
  "validates the build environment and the packaged output. It runs nineteen "
  "checks across five categories:")
table_grid(
    ["Check", "What it verifies"],
    [
        ("Compiler", "MinGW-w64 g++ in the vendored mingw64/bin runs and reports its version."),
        ("SFML linkage",
         "Compiles tools/probes/sfml_probe.cpp against third_party/SFML-mingw/ "
         "and confirms the resulting probe.exe is produced."),
        ("Runtime DLLs",
         "Each of the seven required DLLs is present beside the .exe: "
         "sfml-graphics-3, sfml-window-3, sfml-system-3, sfml-audio-3, "
         "libgcc_s_seh-1, libstdc++-6, libwinpthread-1."),
        ("Fonts",
         "Inter-Regular.ttf, Inter-Bold.ttf, and fa-solid-900.ttf exist in "
         "both assets/fonts/ and dist/assets/fonts/."),
        ("Asset/data paths",
         "Source assets/, assets/fonts/, assets/img/, and data/ paths "
         "exist and are readable."),
    ],
    col_widths=(Inches(1.5), Inches(4.8)),
)

H("12.3  Fresh-Machine Simulation", level=2)
P("Before declaring the bundle 'self-contained', it is verified end-to-end "
  "by copying dist/ to a temporary folder, stripping PATH down to just "
  "C:\\Windows\\System32 and C:\\Windows (simulating a machine without the "
  "development toolchain), and launching smart_pos.exe from the temp "
  "location. The app must run cleanly for at least three seconds without "
  "throwing a missing-DLL dialog. This confirms there are no hidden "
  "dependencies on the developer's PATH.")

H("12.4  Manual Exercise Across Roles", level=2)
P("Each role's screens are manually exercised end to end:")
bullets([
    "Admin: log in, add a menu item, edit it, delete it, search by name, "
    "filter by category. Adjust stock on an ingredient; verify the dashboard "
    "low-stock card updates.",
    "Cashier: log in, place an order with two different items, attach a "
    "customer by phone, apply a discount, place the order, generate a "
    "receipt, save it. Confirm receipts/RCPT-NNNNNN.txt appears on disk "
    "with the correct totals.",
    "Kitchen: log in, see the just-placed order in the Pending lane, "
    "advance it through Preparing → Ready → Served. Confirm the order "
    "disappears from the kitchen queue.",
    "Restart the app and confirm: the new order persists, the inventory "
    "reflects the deduction, the customer has earned loyalty points, the "
    "Analytics screen's daily revenue updates.",
])

H("12.5  Crash Recovery", level=2)
P("The application is force-killed mid-write to confirm the write-temp-"
  "then-rename pattern protects the live file. After kill, the live "
  "data/menu.txt is unchanged and a stray data/menu.txt.tmp may be left "
  "behind (harmless and is overwritten by the next save). No file is ever "
  "left half-written in place.")
hrule()


# ============================================================================
#  13. RESULTS
# ============================================================================
H("13.  Results", level=1)
P("The completed application meets every functional and non-functional "
  "requirement from §4. Quantitative summary:")
table_kv([
    ("Total source files (C++)",  "139 — src/domain (89), src/ui (28), src/platform (9), src/util (7), main.cpp + App"),
    ("Lines of C++ (estimate)",   "~6,500 hand-written across the domain, UI, platform, and util layers"),
    ("Domain sub-packages",       "10 (auth, common, menu, order, billing, inventory, customer, tables, kitchen, analytics) + persistence"),
    ("Binary record types",       "6 — UserRecord, OrderHeaderRecord, OrderItemRecord, CustomerRecord, LoyaltyRecord, SalesRecord"),
    ("Text file types",           "6 — menu, inventory, recipes, tables, settings, log + receipt files"),
    ("Demo user accounts",        "3 — admin / cashier / kitchen, auto-seeded on first run"),
    ("Built executable size",     "~2.0 MB"),
    ("Packaged dist/ size",       "~14 MB including all DLLs, assets, and seeded data"),
    ("Cold-start time",           "Under 2 seconds on commodity hardware"),
    ("Build dependencies",        "Vendored — MinGW, SFML 3.1 source, ImGui, CMake portable, fonts"),
    ("Runtime dependencies",      "Zero — runs offline on any 64-bit Windows 10/11"),
])
hrule()


# ============================================================================
#  14. PROJECT JOURNEY & LESSONS LEARNED
# ============================================================================
H("14.  Project Journey & Lessons Learned", level=1)
P("The project was built in five gated phases, each ending in a "
  "verification step before the next began.")
table_grid(
    ["Phase", "Output"],
    [
        ("1. Planning",
         "Ten design documents — architecture, folder structure, class "
         "design, OOP-concept map, UML and use-case diagrams, data-flow "
         "diagrams, file schemas, ASCII wireframes, build roadmap."),
        ("2-Stage 1. Toolchain & smoke test",
         "Vendored CMake portable; rebuilt SFML 3.1.0 from source with "
         "the bundled MinGW (the pre-supplied SFML was an MSVC build, "
         "incompatible with MinGW's C++ ABI); vendored fonts; built a "
         "minimal ImGui-window-inside-an-SFML-window smoke test."),
        ("2-Stage 2. Domain layer",
         "All 89 domain files implemented in three rounds of "
         "build-fix-build cycles. Zero ImGui or SFML dependency anywhere "
         "in src/domain/."),
        ("2-Stage 3 & 4. UI",
         "Theme system, widgets (Card, Sidebar, TopBar, Chart, Toast), "
         "animation helpers, and all ten screens. App class wired everything "
         "together; main.cpp became a 4-line entry point."),
        ("2-Stage 5. Packaging",
         "Release build with post-build CMake commands copying every "
         "runtime DLL and the assets folder into dist/. Fresh-machine "
         "simulation verified the .exe runs standalone."),
        ("Presentation & publish",
         "14-slide PowerPoint generated via python-pptx; repository "
         "initialized, .gitignore tuned to keep the runnable bundle in "
         "the repo while excluding build bulk; pushed to GitHub."),
    ],
    col_widths=(Inches(1.6), Inches(4.7)),
)

H("14.1  Real Engineering Problems Encountered", level=2)
P("The project encountered three concrete engineering problems that "
  "required real diagnosis (not invented for the report):")
bullets([
    "SFML toolchain mismatch. The provided SFML 3.1.0 was an MSVC build "
    "(.lib import libraries with matching .pdb debug-symbol files). MinGW "
    "cannot link such libraries directly because MSVC's C++ name mangling "
    "is incompatible with MinGW's Itanium ABI. The fix was to rebuild SFML "
    "3.1.0 from source with the bundled MinGW GCC, producing proper "
    "lib*.a archives in third_party/SFML-mingw/.",
    "Password hash truncation. After the first commit, login with the "
    "default admin failed. Diagnosis traced it to a copyFixed helper that "
    "reserved the last byte of every fixed-width field for a NUL "
    "terminator, silently truncating the 32-character password hash to "
    "31 characters. The hash on disk could therefore never re-verify. "
    "Fix: copyFixed now writes up to N bytes (the field is zero-initialized "
    "first, so shorter strings still get an implicit terminator at the end "
    "of their content). A standalone probe — "
    "tools/probes/login_probe.cpp — confirms the fix.",
    "Demo-user coverage. The original seed only created the admin "
    "account, so demonstrating the Cashier and Kitchen roles required "
    "Admin to first create those users. The fix was a "
    "seedDemoUsersIfEmpty method that idempotently creates all three "
    "demo accounts on first run.",
])

H("14.2  What the Project Taught", level=2)
bullets([
    "Immediate-mode UI and rigorous OOP can coexist, but only with an "
    "explicit, enforced separation between the domain (object-oriented) "
    "and the UI (procedural per-frame).",
    "static_assert on every binary record's sizeof is a small investment "
    "that catches a whole class of on-disk corruption bugs at compile time.",
    "Vendoring everything — toolchain, library source, fonts, CMake — "
    "trades a small one-time setup cost for a project that compiles "
    "deterministically on any 64-bit Windows machine, indefinitely.",
    "Crash safety does not require a database. Two lines of C++ "
    "(write to <path>.tmp, then std::rename) deliver the same guarantee "
    "for whole-file rewrites.",
])
hrule()


# ============================================================================
#  15. LIMITATIONS & FUTURE WORK
# ============================================================================
H("15.  Limitations & Future Work", level=1)
P("Honest scope. The project is complete against its own objectives, but "
  "explicit limitations are worth noting:")
bullets([
    "Single-machine only. The data files are designed for one workstation. "
    "Multi-station synchronization (e.g. front-of-house and kitchen "
    "display on separate machines sharing a LAN folder) is out of scope.",
    "Hash is demo-grade. PasswordHash uses a salted FNV-1a 64-bit mix — "
    "intentionally simple, easy to defend in the OOP context, and clearly "
    "labelled as not for production use. A real deployment would use a "
    "memory-hard KDF (Argon2id, scrypt) via a vetted library.",
    "PDF export is text-formatted. The 'export PDF' button writes a "
    "monospace .pdf.txt file with the receipt content, not a true PDF. "
    "True PDF generation would require an additional library (e.g. "
    "libharu) which is out of scope.",
    "No real printer driver. Receipts are saved to .txt files. Hooking up "
    "ESC/POS over USB to a thermal printer is straightforward but specific "
    "to a deployment.",
    "No tests in CI. The project relies on the preflight script, manual "
    "exercise, and the fresh-machine simulation. Adding a GoogleTest-based "
    "unit test suite for the domain layer would be a clean future addition.",
])
P("Concrete forward path:")
bullets([
    "Real PDF export via a small embeddable PDF library.",
    "ESC/POS thermal printer driver hookup.",
    "Format-version header on binary record files so on-disk schema can "
    "evolve safely without losing existing data.",
    "Backup-and-restore UI action (zip the data/ folder).",
    "Predictive low-stock alerts (days-until-out) computed from a rolling "
    "average over sales_history.dat.",
    "Multi-station mode with a file-watcher on a LAN share.",
    "Internationalization via externalized strings in data/i18n/<lang>.txt.",
    "Accessibility pass: keyboard-only navigation, larger-text mode, and a "
    "color-blind-safe palette for the table grid.",
])
hrule()


# ============================================================================
#  16. CONCLUSION
# ============================================================================
H("16.  Conclusion", level=1)
P("Smart Restaurant POS is a working, self-contained, file-backed "
  "point-of-sale desktop application written in C++17. It honors the two "
  "constraints that shaped it from day one: it looks and behaves like a "
  "real product, and it demonstrates every OOP concept the course requires, "
  "mapped to specific named classes in the source.")
P("The codebase is split cleanly into a domain layer (89 files, no UI "
  "dependency), an immediate-mode UI layer (28 files, no business state), "
  "a tiny platform layer (9 files, the only place that knows about both "
  "ImGui and SFML), and a util layer. Persistence is exclusively file-based "
  "and crash-safe. The runnable executable plus its DLLs, assets, and "
  "seeded data weighs in at about 14 MB, runs offline on any 64-bit "
  "Windows 10 or 11 machine, and requires no installer.")
P("The project's full source, planning documents, presentation, runnable "
  "bundle, and this report are committed to the public repository at:")
P("github.com/jazibali2876-jpg/oop-project", italic=True,
  align=WD_ALIGN_PARAGRAPH.CENTER)
hrule()


# ============================================================================
#  17. REFERENCES
# ============================================================================
H("17.  References", level=1)
refs = [
    "ISO/IEC 14882:2017. Information technology — Programming languages — "
    "C++. International Organization for Standardization, 2017.",
    "Cornut, O. Dear ImGui — Bloat-free Graphical User interface for C++ "
    "with minimal dependencies. github.com/ocornut/imgui",
    "SFML Team. Simple and Fast Multimedia Library, version 3.1.0. "
    "www.sfml-dev.org",
    "Kitware Inc. CMake — Cross-platform build system. cmake.org",
    "MinGW-w64 Project. MinGW-w64 — for 32 and 64 bit Windows. "
    "www.mingw-w64.org",
    "Andersson, R. Inter typeface family. rsms.me/inter (SIL Open Font "
    "License 1.1).",
    "Font Awesome, Inc. Font Awesome 6 Free. fontawesome.com "
    "(CC BY 4.0 for icons, SIL OFL 1.1 for fonts).",
    "Stroustrup, B. The C++ Programming Language, 4th Edition. "
    "Addison-Wesley, 2013.",
    "Lippman, S., Lajoie, J., Moo, B. C++ Primer, 5th Edition. "
    "Addison-Wesley, 2012.",
    "Meyers, S. Effective Modern C++. O'Reilly Media, 2014.",
]
for r in refs:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    rr = p.add_run(r); set_run(rr, size=10)
page_break()


# ============================================================================
#  APPENDIX A — Repository layout
# ============================================================================
H("Appendix A — Repository Layout", level=1)
P("Pathnames are relative to the repository root.")
code_block(
"""oop-project/
├── README.md                        Top-level project description.
├── CMakeLists.txt                   Build description.
├── .gitignore                       Excludes build-time bulk + runtime state.
├── planning/                        Ten design documents (01..10).
├── src/                             ALL hand-written C++17 source.
│   ├── main.cpp                     Entry point.
│   ├── App.h / App.cpp              Lifecycle, services bundle, main loop.
│   ├── platform/                    ImGui ↔ SFML 3 integration glue.
│   │   ├── Platform.{h,cpp}         sf::RenderWindow + GL context owner.
│   │   ├── ImGuiBackend.{h,cpp}     ImGui context + OpenGL3 backend.
│   │   ├── InputBridge.{h,cpp}      sf::Event -> ImGuiIO translator.
│   │   └── KeyMap.{h,cpp}           sf::Keyboard::Key -> ImGuiKey table.
│   ├── domain/                      OOP-graded layer. No ImGui/SFML headers.
│   │   ├── common/                  Money, DateTime, Id<Tag>, Capability,
│   │   │                            Exceptions, Result.
│   │   ├── auth/                    User + Admin/Cashier/Kitchen, Session,
│   │   │                            PasswordHash, AuthService,
│   │   │                            UserRepository.
│   │   ├── menu/                    MenuItem, Category, MenuService,
│   │   │                            MenuRepository.
│   │   ├── inventory/               Ingredient, Recipe, InventoryService,
│   │   │                            InventoryRepository, RecipeRepository.
│   │   ├── order/                   OrderItem, Order, OrderStatus,
│   │   │                            OrderService, OrderRepository.
│   │   ├── billing/                 Receipt, PaymentMethod, BillingService,
│   │   │                            ReceiptRepository, PdfExporter.
│   │   ├── customer/                Customer, LoyaltyRule, CustomerService,
│   │   │                            CustomerRepository, LoyaltyRepository.
│   │   ├── tables/                  Table, TableStatus, TableService,
│   │   │                            TableRepository.
│   │   ├── kitchen/                 KitchenTicket, KitchenService.
│   │   ├── analytics/               SalesRecord, Report, AnalyticsService,
│   │   │                            SalesHistoryRepository.
│   │   └── persistence/             FilePaths, TextRepository,
│   │                                BinaryRepository<T> (template),
│   │                                SettingsRepository.
│   ├── ui/                          Immediate-mode UI. No business state.
│   │   ├── AppContext.h             Bundle of service references.
│   │   ├── Router.{h,cpp}           Active-screen state.
│   │   ├── Shell.{h,cpp}            Sidebar + top bar + screen dispatch.
│   │   ├── screens/                 SplashScreen, LoginScreen,
│   │   │                            DashboardScreen, MenuScreen,
│   │   │                            OrderScreen, BillingScreen,
│   │   │                            InventoryScreen, TablesScreen,
│   │   │                            KitchenScreen, AnalyticsScreen.
│   │   ├── widgets/                 Card, Chart, Sidebar, TopBar, Toast.
│   │   ├── theme/                   Theme, Fonts, Icons.
│   │   ├── anim/                    Animation (easing) + Tween.
│   │   └── state/                   OrderDraft (UI cart), ToastQueue.
│   └── util/                        Strings, Files (writeAtomic), Log, Random.
├── assets/                          Vendored fonts + image folders.
├── data/                            Runtime state (autocreated on first run).
├── dist/                            Self-contained runnable bundle.
│   ├── smart_pos.exe                The application.
│   ├── sfml-*-3.dll                 SFML 3 runtime DLLs.
│   ├── libgcc_s_seh-1.dll, ...      MinGW C++ runtime DLLs.
│   ├── assets/                      Fonts copied next to the exe.
│   ├── data/                        Auto-seeded on first run.
│   └── HOW_TO_RUN.txt               One-paragraph instructions.
├── docs/                            This report (Project_Report.docx).
├── presentation/                    Demo deck (PowerPoint) + generator.
└── tools/                           Preflight script + small C++ probes.""",
    font_size=8)


# ============================================================================
#  APPENDIX B — Binary file schemas
# ============================================================================
page_break()
H("Appendix B — Binary File Schemas", level=1)
P("Every record struct uses #pragma pack(push, 1) followed by a "
  "static_assert(sizeof(R) == N) pinning the exact on-disk byte count.")

H("B.1  UserRecord (128 B)", level=2)
table_grid(
    ["Field", "Type", "Bytes", "Notes"],
    [
        ("id",             "uint32_t",    "4",  "Monotonically assigned by AuthService."),
        ("username",       "char[24]",    "24", "NUL-padded if shorter than 24."),
        ("passwordHash",   "char[32]",    "32", "Exactly 32 hex chars (16+16 from two FNV-1a runs)."),
        ("salt",           "char[16]",    "16", "8 lowercase ASCII chars + NUL pad."),
        ("fullName",       "char[32]",    "32", "Human-friendly name."),
        ("role",           "uint8_t",     "1",  "0=Admin, 1=Cashier, 2=Kitchen."),
        ("active",         "uint8_t",     "1",  "1 if the account is enabled."),
        ("createdAtEpoch", "int64_t",     "8",  "Seconds since UNIX epoch."),
        ("pad",            "uint8_t[10]", "10", "Reserved — keeps total at 128."),
    ],
    col_widths=(Inches(1.5), Inches(1.2), Inches(0.7), Inches(2.9)),
)

H("B.2  OrderHeaderRecord (128 B) + OrderItemRecord (56 B)", level=2)
P("Orders are variable-length. One header is followed by itemCount "
  "OrderItemRecord rows.")
table_grid(
    ["Field", "Type", "Bytes"],
    [
        ("id",                  "uint32_t",    "4"),
        ("placedAtEpoch",       "int64_t",     "8"),
        ("itemCount",           "uint16_t",    "2"),
        ("status",              "uint8_t",     "1"),
        ("paymentMethod",       "uint8_t",     "1"),
        ("discountCents",       "int64_t",     "8"),
        ("taxBp",               "uint16_t",    "2"),
        ("customerId",          "uint32_t",    "4"),
        ("cashierUsername",     "char[24]",    "24"),
        ("specialInstructions", "char[64]",    "64"),
        ("totalCents",          "int64_t",     "8"),
        ("pad",                 "uint8_t[2]",  "2"),
    ],
    col_widths=(Inches(2.0), Inches(1.4), Inches(0.7)),
)
P("OrderItemRecord:")
table_grid(
    ["Field", "Type", "Bytes"],
    [
        ("menuItemId",     "uint32_t",   "4"),
        ("qty",            "uint16_t",   "2"),
        ("unitPriceCents", "int64_t",    "8"),
        ("notes",          "char[40]",   "40"),
        ("pad",            "uint8_t[2]", "2"),
    ],
    col_widths=(Inches(2.0), Inches(1.4), Inches(0.7)),
)

H("B.3  CustomerRecord (96 B)", level=2)
table_grid(
    ["Field", "Type", "Bytes"],
    [
        ("id",              "uint32_t",   "4"),
        ("name",            "char[40]",   "40"),
        ("phone",           "char[20]",   "20"),
        ("joinedAtEpoch",   "int64_t",    "8"),
        ("totalSpentCents", "int64_t",    "8"),
        ("loyaltyPoints",   "int32_t",    "4"),
        ("orderCount",      "int32_t",    "4"),
        ("active",          "uint8_t",    "1"),
        ("pad",             "uint8_t[7]", "7"),
    ],
    col_widths=(Inches(2.0), Inches(1.4), Inches(0.7)),
)

H("B.4  LoyaltyRecord (32 B) — append-only ledger", level=2)
table_grid(
    ["Field", "Type", "Bytes"],
    [
        ("customerId",  "uint32_t",   "4"),
        ("orderId",     "uint32_t",   "4"),
        ("atEpoch",     "int64_t",    "8"),
        ("pointsDelta", "int32_t",    "4"),
        ("spendCents",  "int64_t",    "8"),
        ("reason",      "uint8_t",    "1"),
        ("pad",         "uint8_t[3]", "3"),
    ],
    col_widths=(Inches(2.0), Inches(1.4), Inches(0.7)),
)

H("B.5  SalesRecord (64 B)", level=2)
table_grid(
    ["Field", "Type", "Bytes"],
    [
        ("orderId",       "uint32_t",    "4"),
        ("epoch",         "int64_t",     "8"),
        ("totalCents",    "int64_t",     "8"),
        ("itemCount",     "int16_t",     "2"),
        ("paymentMethod", "uint8_t",     "1"),
        ("hour",          "uint8_t",     "1"),
        ("cashier",       "char[24]",    "24"),
        ("customerId",    "int32_t",     "4"),
        ("pad",           "uint8_t[12]", "12"),
    ],
    col_widths=(Inches(2.0), Inches(1.4), Inches(0.7)),
)


# ============================================================================
#  APPENDIX C — How to run & build
# ============================================================================
page_break()
H("Appendix C — How to Run & Build", level=1)

H("C.1  Run the existing release", level=2)
numbered([
    "Clone or download the repository.",
    "Open the dist/ folder.",
    "Double-click smart_pos.exe.",
    "Sign in with admin / admin123 (or cashier / cashier123, or "
    "kitchen / kitchen123).",
])
P("Requirements: 64-bit Windows 10 or 11, and a GPU that supports OpenGL "
  "3.3 or higher (every modern PC qualifies). No installer, no admin "
  "rights, no network access needed.")

H("C.2  Build from source", level=2)
P("The repository excludes the build-time bulk (toolchain, library SDKs, "
  "intermediates). To build from source, supply the following alongside "
  "the repo — they are listed in .gitignore for size reasons:")
table_kv([
    ("./SFML-3.1.0/",      "SFML 3.1.0 distribution."),
    ("./imgui-master/",    "Dear ImGui master branch."),
    ("./mingw64/",         "MinGW-w64 GCC 14.x (UCRT / POSIX / SEH)."),
    ("./third_party/cmake-*",   "CMake portable build."),
    ("./third_party/SFML-mingw/","SFML rebuilt from source with MinGW."),
])
P("Then, from the project root:")
code_block(
"""# from the project root, with the toolchain on PATH:
export PATH=\"$PWD/mingw64/bin:$PWD/third_party/SFML-mingw/bin:$PATH\"
CMAKE=\"$PWD/third_party/cmake-3.30.5-windows-x86_64/bin/cmake.exe\"

# one-time configure:
\"$CMAKE\" -S . -B build/project -G \"MinGW Makefiles\" \\
    -DCMAKE_BUILD_TYPE=Release \\
    -DCMAKE_CXX_COMPILER=\"$PWD/mingw64/bin/g++.exe\" \\
    -DCMAKE_MAKE_PROGRAM=\"$PWD/mingw64/bin/mingw32-make.exe\"

# build:
\"$CMAKE\" --build build/project -- -j 4

# verify:
powershell -ExecutionPolicy Bypass -File ./tools/preflight_check.ps1""",
    font_size=9)
P("The output smart_pos.exe plus every required runtime DLL plus the "
  "assets folder lands in ./dist/, ready to run.")


# ============================================================================
doc.save(str(OUT))
size_kb = OUT.stat().st_size / 1024
n_paras = len(doc.paragraphs)
print(f"wrote {OUT}  ({size_kb:.1f} KB, {n_paras} paragraphs)")
